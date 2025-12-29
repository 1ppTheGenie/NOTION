"""
Chargeback Response Generator - Word Document Version
Creates a .docx file that can be manually edited
"""
import json
import sys
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

# Color Palette
COLORS = {
    'primary': RGBColor(44, 62, 80),      # #2c3e50
    'secondary': RGBColor(52, 152, 219),   # #3498db
    'success': RGBColor(39, 174, 96),     # #27ae60
    'dark': RGBColor(52, 73, 94),          # #34495e
}

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_table_style(table, header_color=None):
    """Apply styling to table"""
    if header_color:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), header_color)
        table.rows[0].cells[0]._element.get_or_add_tcPr().append(shading_elm)
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

def generate_chargeback_response_docx(evidence_file, kit_dir, terms_file=None):
    """Generate Word document version of chargeback response"""
    print("\n" + "="*80)
    print("GENERATING CHARGEBACK RESPONSE - WORD DOCUMENT VERSION")
    print("="*80)
    
    # Load evidence
    with open(evidence_file, 'r', encoding='utf-8') as f:
        evidence = json.load(f)
    
    case = evidence['case_info']
    transaction_id = case['paypal_transaction_id']
    
    # Get cardholder name
    cardholder_name = "Customer"
    if evidence.get('user_details'):
        user = evidence['user_details']
        if user.get('UserName'):
            cardholder_name = user['UserName']
        elif user.get('Email'):
            cardholder_name = user['Email'].split('@')[0]
    
    # Clean name
    cardholder_name_clean = re.sub(r'[^\w\s-]', '', cardholder_name)
    cardholder_name_clean = re.sub(r'[-\s]+', '_', cardholder_name_clean).strip('_')
    
    # Date formatting
    trans_date_obj = datetime.strptime(case.get('transaction_date', datetime.now().strftime('%Y-%m-%d')), '%Y-%m-%d')
    file_date = trans_date_obj.strftime('%m_%d_%Y')
    
    # Version tracking
    base_filename = f"Response_to_{cardholder_name_clean}_Dispute_{file_date}"
    version = 1
    existing_files = list(Path(kit_dir).glob(f"{base_filename}_v*.docx"))
    if existing_files:
        versions = []
        for f in existing_files:
            match = re.search(r'_v(\d+)\.docx$', f.name)
            if match:
                versions.append(int(match.group(1)))
        if versions:
            version = max(versions) + 1
    
    # Filename
    docx_filename = f"{base_filename}_v{version}.docx"
    docx_path = Path(kit_dir) / docx_filename
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========================================================================
    # COVER PAGE - PAGE 1
    # ========================================================================
    # Title
    title = doc.add_heading('CHARGEBACK RESPONSE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.color.rgb = COLORS['primary']
    
    # Subtitle
    subtitle = doc.add_paragraph(f'Response to Dispute Filed by {cardholder_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = COLORS['dark']
    
    doc.add_paragraph()  # Spacing
    
    # Document metadata table
    now = datetime.now()
    meta_table = doc.add_table(rows=7, cols=2)
    meta_table.style = 'Light Grid Accent 1'
    
    meta_data = [
        ['Document Name', 'Chargeback Response - Dispute Resolution'],
        ['Version', f'{version}.0'],
        ['Date Created', now.strftime('%m/%d/%Y')],
        ['Time Created', now.strftime('%I:%M %p')],
        ['Created By', 'TheGenie.ai Customer Experience Team'],
        ['Email', 'wecare@thegenie.ai'],
        ['Phone', '888-425-2300']
    ]
    
    for i, (label, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        meta_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================================================
    # TABLE OF CONTENTS - PAGE 2
    # ========================================================================
    toc_heading = doc.add_heading('TABLE OF CONTENTS', 1)
    toc_heading.runs[0].font.color.rgb = COLORS['primary']
    
    doc.add_paragraph()  # Spacing
    
    # TOC items with page numbers
    toc_items = [
        ("1. PayPal Chargeback Case Details", 3),
        ("2. What Was Ordered", 4),
        ("3. Customer Resolution Attempt", 6),
        ("4. The Story: What Actually Happened", 7),
        ("5. Workflow Screenshots", 8),
        ("6. Evidence Summary", 10),
        ("7. Proof of Authorization", 11),
        ("8. Proof of Service Delivery", 12),
        ("9. Proof of No Contact", 13),
        ("10. Terms of Service", 14),
        ("11. Conclusion & Request", 15)
    ]
    
    # Create TOC table
    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.columns[0].width = Inches(4.5)
    toc_table.columns[1].width = Inches(1.5)
    toc_table.columns[2].width = Inches(0.5)
    
    for i, (item, page_num) in enumerate(toc_items):
        toc_table.rows[i].cells[0].text = item
        toc_table.rows[i].cells[1].text = '.' * 50  # Dots
        toc_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(128, 128, 128)
        toc_table.rows[i].cells[2].text = str(page_num)
        toc_table.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # ========================================================================
    # PAYPAL CHARGEBACK CASE DETAILS - PAGE 3
    # ========================================================================
    case_heading = doc.add_heading('PAYPAL CHARGEBACK CASE DETAILS', 1)
    case_heading.runs[0].font.color.rgb = COLORS['primary']
    
    doc.add_paragraph('The following information is from the PayPal Resolution Center for this chargeback case:')
    doc.add_paragraph()  # Spacing
    
    # Case details table
    case_table = doc.add_table(rows=7, cols=2)
    case_table.style = 'Light Grid Accent 1'
    case_table.columns[0].width = Inches(2.5)
    case_table.columns[1].width = Inches(3.5)
    
    case_data = [
        ['Field', 'Value'],
        ['Case ID', transaction_id],
        ['Transaction Amount', f"${case.get('transaction_amount', '67.50')} USD"],
        ['Disputed Amount', f"${case.get('transaction_amount', '67.50')} USD"],
        ['Transaction Date', trans_date_obj.strftime('%B %d, %Y')],
        ['Buyer Name', cardholder_name],
        ['Chargeback Reason', 'The buyer stated that they did not make this purchase.']
    ]
    
    for i, (label, value) in enumerate(case_data):
        case_table.rows[i].cells[0].text = label
        case_table.rows[i].cells[1].text = value
        if i == 0:  # Header row
            for cell in case_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '3498db')
                cell._element.get_or_add_tcPr().append(shading)
    
    doc.add_paragraph()  # Spacing
    
    # Response to chargeback reason
    response_para = doc.add_paragraph()
    response_para.add_run('Response to Chargeback Reason: ').bold = True
    response_para.add_run('The buyer\'s claim that they "did not make this purchase" is directly contradicted by the evidence presented in this document. The evidence demonstrates:')
    
    doc.add_paragraph('• The purchase was made from the buyer\'s verified account')
    doc.add_paragraph('• The buyer actively used the service on the transaction date')
    doc.add_paragraph('• The buyer\'s account shows clear activity logs demonstrating service usage')
    doc.add_paragraph('• All authentication and authorization evidence confirms the buyer\'s identity')
    doc.add_paragraph('• The service was fully delivered and executed as ordered')
    
    doc.add_paragraph('This document provides comprehensive evidence refuting the buyer\'s claim and demonstrates that the transaction was legitimate, authorized, and the service was fully delivered.')
    doc.add_paragraph()  # Spacing
    
    # Transaction info table
    trans_table = doc.add_table(rows=5, cols=2)
    trans_table.style = 'Light Grid Accent 1'
    trans_table.columns[0].width = Inches(2.5)
    trans_table.columns[1].width = Inches(3.5)
    
    trans_data = [
        ['Transaction ID', transaction_id],
        ['Customer', case['customer_email']],
        ['Transaction Date', case['transaction_date']],
        ['Amount', f"${case['transaction_amount']}"],
        ['Product', 'Listing Command - Digital Service']
    ]
    
    for i, (label, value) in enumerate(trans_data):
        trans_table.rows[i].cells[0].text = label
        trans_table.rows[i].cells[1].text = value
        trans_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        if i == 0:  # Header row
            for cell in trans_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '2c3e50')
                cell._element.get_or_add_tcPr().append(shading)
    
    doc.add_paragraph()  # Spacing
    
    # ========================================================================
    # WHAT WAS ORDERED - Continue on page 3
    # ========================================================================
    order_heading = doc.add_heading('WHAT WAS ORDERED', 1)
    order_heading.runs[0].font.color.rgb = COLORS['primary']
    
    # Narrative
    narrative = doc.add_paragraph()
    narrative.add_run('Facts. Not opinions. Not claims. Facts.\n\n').bold = True
    narrative.add_run(f'On {case.get("transaction_date", "December 5, 2025")}, {cardholder_name} executed a transaction. ')
    narrative.add_run('Not a mistake. Not an accident. A deliberate, authenticated, verified purchase of ').bold = True
    narrative.add_run('Listing Command')
    narrative.add_run(' - a one-time use digital marketing service that was immediately executed and delivered.\n\n')
    narrative.add_run('This isn\'t about what someone says happened. This is about what ').italic = True
    narrative.add_run('actually')
    narrative.add_run(' happened. The evidence doesn\'t lie. The data doesn\'t have an agenda. The logs don\'t have feelings. They simply record reality.')
    
    doc.add_paragraph()  # Spacing
    
    # Order Summary heading
    summary_para = doc.add_paragraph()
    summary_para.add_run('Order Summary').bold = True
    
    # Order Details Table - This is the table that needs to fit on one page
    order_table = doc.add_table(rows=11, cols=2)
    order_table.style = 'Light Grid Accent 1'
    order_table.columns[0].width = Inches(2.5)
    order_table.columns[1].width = Inches(3.5)
    
    order_details = [
        ['Property Address', '1816 9th Street, Manhattan Beach, CA 90266'],
        ['MLS Number', 'SB25228445'],
        ['Area', 'East Manhattan Beach'],
        ['Service Type', 'SMS Text Messaging Campaign'],
        ['Target Audience Size', '150 Properties'],
        ['Messages Delivered', '149'],
        ['Engagements Received', '1'],
        ['Collection ID', '1c7bdd67-9701-4159-8fa7-4f4a26c5e432'],
        ['Order Date', case.get('transaction_date', '2025-12-05')],
        ['Processing Date', 'December 5, 2025']
    ]
    
    # Header row
    order_table.rows[0].cells[0].text = 'Field'
    order_table.rows[0].cells[1].text = 'Value'
    for cell in order_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '2c3e50')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Data rows
    for i, (label, value) in enumerate(order_details, 1):
        order_table.rows[i].cells[0].text = label
        order_table.rows[i].cells[1].text = value
    
    # Set smaller font for table to help it fit
    for row in order_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    # Save document
    doc.save(str(docx_path))
    
    file_size_mb = docx_path.stat().st_size / (1024 * 1024)
    
    print(f"\n✅ Word Document Generated: {docx_filename}")
    print(f"   Cardholder: {cardholder_name}")
    print(f"   File Size: {file_size_mb:.2f} MB")
    print(f"\nDOCX Location: {docx_path}")
    print("\nYou can now open this Word document and manually adjust the table formatting.")
    
    return str(docx_path)

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python generate_chargeback_response_docx.py <evidence_file> <kit_dir> [terms_file]")
        sys.exit(1)
    
    evidence_file = sys.argv[1]
    kit_dir = sys.argv[2]
    terms_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    generate_chargeback_response_docx(evidence_file, kit_dir, terms_file)

